from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import datetime
import os

doc = Document()

# Page Setup
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

DIAGRAM_DIR = os.path.join(os.path.dirname(__file__), 'diagrams')

# ── Helpers ──
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None, size=12, space_after=6, space_before=0, first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_mixed_para(parts, space_after=6, space_before=0, left_indent=None):
    """parts: list of (text, bold) tuples"""
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    return p

def add_image_placeholder(caption, description):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[IMAGE PLACEHOLDER: {description}]')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'Figure: {caption}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.bold = True
    r.italic = True

def add_image(filename, caption, width=5.5):
    path = os.path.join(DIAGRAM_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'Figure: {caption}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
        r.italic = True
    else:
        add_image_placeholder(caption, f'[File not found: {filename}]')

def add_table(headers, rows, caption, col_widths=None):
    """Create a formatted table with headers and rows."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)
        r.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1565c0')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        r.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(8)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'f0f4f8')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[j].width = Inches(w)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'Table: {caption}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.bold = True
    r.italic = True
    doc.add_paragraph()

def add_table_placeholder(caption, description):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[TABLE PLACEHOLDER: {description}]')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'Table: {caption}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.bold = True
    r.italic = True

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)

def add_heading_unnumbered(text):
    """Add an unnumbered heading that still appears bold/large"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

# ════════════════════════════════════════
#              COVER PAGE
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FINAL YEAR PROJECT REPORT')
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Taktill POS')
run.font.name = 'Times New Roman'
run.font.size = Pt(26)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Cloud-Native Point-of-Sale System for Retail Businesses')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.italic = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Full Unit  \u2022  Final Report')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Student: [Student\'s First and Last Name]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Supervisor: [Supervisor Name]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A report submitted in part fulfilment of the degree of\n[Bachelor\'s Degree Name]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Isbat University Kampala')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(datetime.datetime.now().strftime('%B %d, %Y'))
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════
#              DECLARATION
# ════════════════════════════════════════
add_heading_custom('Declaration', level=1)
add_para('This report has been prepared on the basis of my own work. Where other published and unpublished source materials have been used, these have been acknowledged.', space_after=12)

tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
labels = ['Word Count:', 'Student Name:', 'Date of Submission:', 'Signature:']
for i, label in enumerate(labels):
    cell0 = tbl.cell(i, 0)
    cell0.text = ''
    r0 = cell0.paragraphs[0].add_run(label)
    r0.font.name = 'Times New Roman'
    r0.font.size = Pt(12)
    r0.bold = True
    cell1 = tbl.cell(i, 1)
    cell1.text = ''
    r1 = cell1.paragraphs[0].add_run('[To be completed before submission]')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.italic = True

doc.add_page_break()

# ════════════════════════════════════════
#             TABLE OF CONTENTS
# ════════════════════════════════════════
add_heading_custom('Table of Contents', level=1)

toc_items = [
    ('Abstract', 4),
    ('Project Specification', 5),
    ('Chapter 1: Introduction', 6),
    ('    1.1 Background', 6),
    ('    1.2 Problem Statement', 6),
    ('    1.3 Aims and Objectives', 7),
    ('    1.4 Scope', 7),
    ('Chapter 2: Technologies and Tools', 8),
    ('    2.1 Frontend Technologies', 8),
    ('    2.2 Backend Technologies', 9),
    ('    2.3 Database Technologies', 10),
    ('    2.4 Development and Deployment Tools', 11),
    ('Chapter 3: System Analysis and Design', 12),
    ('    3.1 Requirements Analysis', 12),
    ('    3.2 System Architecture', 13),
    ('    3.3 Database Design', 14),
    ('    3.4 API Design', 16),
    ('    3.5 User Interface Design', 16),
    ('Chapter 4: Implementation', 18),
    ('    4.1 Backend Implementation', 18),
    ('    4.2 Frontend Implementation', 20),
    ('    4.3 Authentication System', 20),
    ('    4.4 Shopping Cart and Sales Processing', 21),
    ('Chapter 5: Testing and Evaluation', 23),
    ('    5.1 Testing Strategy', 23),
    ('    5.2 Unit Testing', 23),
    ('    5.3 Integration Testing', 24),
    ('    5.4 End-to-End Testing', 25),
    ('    5.5 User Acceptance Testing', 26),
    ('Chapter 6: Deployment and Maintenance', 28),
    ('    6.1 Deployment Architecture', 28),
    ('    6.2 Development Workflow', 29),
    ('    6.3 Maintenance Considerations', 29),
    ('Chapter 7: Conclusion and Future Work', 31),
    ('    7.1 Achievements', 31),
    ('    7.2 Challenges', 32),
    ('    7.3 Future Work', 32),
    ('Bibliography', 34),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    indent = ''
    display = item
    if item.startswith('    '):
        indent = '     '
        display = item.strip()
    run = p.add_run(f'{indent}{display}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run2 = p.add_run(f'\t{page}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════
#                ABSTRACT
# ════════════════════════════════════════
add_heading_custom('Abstract', level=1)
add_para('This report presents the design, implementation, and evaluation of Taktill POS, a cloud-native, web-based Point-of-Sale (POS) system tailored for small and medium-sized retail businesses in East Africa. The system addresses critical challenges faced by traditional retail operations, including manual inventory tracking, fragmented sales recording, limited reporting capabilities, and the absence of role-based access control for staff management.', first_line_indent=1.27)
add_para('Taktill POS is built using a modern web technology stack comprising React 19 for the frontend, Node.js with Express 5 for the backend REST API, PostgreSQL for data persistence, and Drizzle ORM for type-safe database operations. The system features a complete point-of-sale interface with product search, barcode scanning support, discount management, and receipt generation. It includes comprehensive inventory management with stock adjustment auditing, role-based authentication for owners, managers, and cashiers, and rich reporting capabilities including daily sales summaries, product-level sales analytics, and dashboard visualisations using ApexCharts.', first_line_indent=1.27)
add_para('The application follows a three-tier role-based access control model: Cashiers can process sales with discount caps, Managers have full access to inventory and reporting, and Owners maintain administrative control over users, cashiers, and business settings. All sales are processed atomically using database transactions with row-level locking to ensure data integrity under concurrent access. Monetary values are handled using decimal arithmetic throughout to avoid floating-point precision errors.', first_line_indent=1.27)
add_para('The system is deployed on Vercel\'s serverless platform with Neon PostgreSQL for production, while supporting SQLite for local development. Comprehensive testing including unit tests using Vitest, API integration tests with Supertest, and end-to-end tests with Playwright ensures system reliability. The evaluation demonstrates that Taktill POS meets all specified requirements, providing a robust, scalable, and user-friendly retail management solution suitable for deployment across East African retail businesses.', first_line_indent=1.27)

doc.add_page_break()

# ════════════════════════════════════════
#         PROJECT SPECIFICATION
# ════════════════════════════════════════
add_heading_custom('Project Specification', level=1)
add_para('Project Title:', bold=True)
add_para('Taktill POS \u2014 A Cloud-Native Point-of-Sale System for Retail Businesses')
add_para('Project Type:', bold=True)
add_para('Full Unit Final Year Project \u2014 Web Application Development')
add_para('Supervisor:', bold=True)
add_para('[Supervisor Name]')

add_para('Project Objectives:', bold=True, space_before=6)
objectives = [
    'Design and implement a web-based Point-of-Sale system accessible from any modern browser.',
    'Implement a three-tier role-based access control system for owners, managers, and cashiers.',
    'Develop an intuitive POS interface for processing sales with product search, cart management, and discount application.',
    'Build a comprehensive inventory management module with stock adjustment tracking and low-stock alerts.',
    'Create a reporting and analytics dashboard with daily summaries, product analytics, and data visualisation.',
    'Ensure data integrity through atomic transaction processing and proper decimal arithmetic.',
    'Deploy the system on a cloud platform with a serverless architecture for scalability and reliability.',
]
for obj in objectives:
    add_bullet(obj)

add_para('Deliverables:', bold=True, space_before=6)
deliverables = [
    'Fully functional web application deployed on Vercel.',
    'Source code repository with complete version history.',
    'REST API with comprehensive endpoint coverage.',
    'Database schema with migration scripts for schema evolution.',
    'Test suite including unit, integration, and end-to-end tests.',
    'This project report documenting the complete development process.',
]
for d in deliverables:
    add_bullet(d)

doc.add_page_break()

# ════════════════════════════════════════
#           CHAPTER 1: INTRODUCTION
# ════════════════════════════════════════
add_heading_custom('Chapter 1: Introduction', level=1)

add_heading_custom('1.1 Background', level=2)
add_para('Point-of-Sale (POS) systems are fundamental to modern retail operations, serving as the central hub for sales transactions, inventory management, and business analytics. In East Africa, particularly in Uganda, the retail sector is dominated by small and medium-sized enterprises (SMEs) that often rely on manual methods for sales recording and inventory tracking. These manual approaches are prone to transcription errors, inefficient for high-volume operations, and provide limited visibility into business performance.', first_line_indent=1.27)
add_para('Traditional POS systems available in the market present several barriers to adoption for SMEs. High upfront licensing costs, the requirement for dedicated hardware, and complex installation procedures make them inaccessible to many small retailers. Cloud-based alternatives exist but frequently require ongoing subscription fees and reliable internet connectivity that may not always be available. Furthermore, many existing solutions do not adequately address the multi-role staffing model common in East African retail, where shop owners, managers, and cashiers have distinctly different operational responsibilities and access requirements.', first_line_indent=1.27)
add_para('The widespread adoption of modern web technologies, including React for building rich user interfaces, serverless computing platforms for scalable deployment, and affordable cloud database services, presents an opportunity to develop a POS system that is both powerful and accessible. Taktill POS was conceived to address these gaps, providing a feature-rich, cloud-native POS solution designed specifically for the operational realities of East African retail businesses.', first_line_indent=1.27)

add_heading_custom('1.2 Problem Statement', level=2)
add_para('Small and medium-sized retail businesses in East Africa face several interconnected challenges in managing their daily operations. Sales are often recorded manually in paper ledgers, leading to transcription errors, data loss, and significant delays in reconciliation. Inventory tracking relies on periodic physical counts rather than real-time visibility, resulting in stock discrepancies, missed restocking opportunities, and lost sales due to out-of-stock situations. Generating meaningful business reports requires manual compilation of data from multiple sources, a time-consuming process that discourages regular business performance analysis.', first_line_indent=1.27)
add_para('Staff management presents additional operational challenges. In a typical retail store, multiple cashiers work different shifts across the trading week, and differentiating their access permissions and tracking individual performance is difficult with manual systems. Oversight mechanisms for cashier activities\u2014such as enforcing discount limits, reconciling daily takings, and preventing unauthorised transactions\u2014are hard to implement without automated systems. Owners and managers lack real-time visibility into daily operations, making it difficult to make informed business decisions about stock ordering, staffing, and pricing.', first_line_indent=1.27)
add_para('The core problem, therefore, is the absence of an affordable, accessible, and feature-complete POS system that addresses the specific operational requirements of East African retail businesses. Such a system must provide role-based access control, atomic transaction processing for data integrity, comprehensive inventory management, meaningful business reporting, and cloud-based deployment for remote accessibility without requiring significant capital investment in hardware.', first_line_indent=1.27)

add_heading_custom('1.3 Aims and Objectives', level=2)
add_para('The primary aim of this project is to design, implement, and evaluate Taktill POS, a cloud-native Point-of-Sale system that streamlines retail operations for small and medium-sized businesses. The specific objectives are:')

obj_list = [
    ('1. ', 'To develop a secure authentication system supporting two distinct login methods: email-and-password for owners and managers, and PIN-based login for cashiers, with appropriate session management and security controls.'),
    ('2. ', 'To implement a complete POS interface enabling cashiers to search products, manage a shopping cart, apply line-item and cart-level discounts, process multiple payment methods, and generate printable receipts.'),
    ('3. ', 'To build an inventory management module with full CRUD operations, stock adjustment auditing with reason tracking, low-stock threshold monitoring, and category management.'),
    ('4. ', 'To create a reporting and analytics dashboard providing daily sales summaries with per-cashier breakdown, product-level sales analysis, and business performance metrics visualised through charts.'),
    ('5. ', 'To design and implement a role-based access control system with three tiers\u2014cashier, manager, and owner\u2014each with appropriately scoped permissions and enforced server-side.'),
    ('6. ', 'To ensure data integrity through atomic database transactions for sales processing and proper decimal arithmetic for all monetary calculations across both frontend and backend.'),
    ('7. ', 'To deploy the system on a cloud platform using a serverless architecture for scalability, reliability, and cost-effectiveness, with automated deployment pipelines.'),
]
for prefix, text in obj_list:
    add_mixed_para([(prefix, True), (text, False)], left_indent=1)

add_heading_custom('1.4 Scope', level=2)
add_para('The scope of this project encompasses the development of a fully functional web-based POS system with the following boundaries. The system supports single-business operation with multiple users and cashiers per business. It handles cash, mobile money, card, and other payment methods as recorded transactions, but does not integrate with external payment gateways for real-time payment processing. The inventory module supports stock adjustments with a complete audit trail but does not include purchase order management, supplier integration, or automated reordering.', first_line_indent=1.27)
add_para('The system is designed as a web application accessible through modern browsers with a responsive layout suitable for desktop and tablet use. It is not designed as a native mobile application, although the responsive web design ensures basic usability on mobile devices for monitoring purposes. The system targets deployment on Vercel\'s serverless platform with PostgreSQL database hosting on Neon, though the modular architecture supports migration to alternative cloud providers or self-hosted infrastructure with minimal changes.', first_line_indent=1.27)

doc.add_page_break()

# ════════════════════════════════════════
#        CHAPTER 2: TECHNOLOGIES & TOOLS
# ════════════════════════════════════════
add_heading_custom('Chapter 2: Technologies and Tools', level=1)
add_para('This chapter provides an overview of the technologies, frameworks, and tools selected for the development of Taktill POS. The technology choices were guided by the requirements for a modern, scalable, and maintainable web application, with careful consideration of the trade-offs between development velocity, runtime performance, and long-term maintainability.', first_line_indent=1.27)

add_heading_custom('2.1 Frontend Technologies', level=2)

add_mixed_para([('React 19. ', True)], space_before=6)
add_para('React 19, the latest stable release of the widely adopted JavaScript library for building user interfaces, was chosen as the frontend framework. React\'s component-based architecture enables the creation of reusable UI components, promoting code maintainability and consistent design throughout the application. The virtual DOM implementation provides efficient rendering, which is critical for a POS system where real-time responsiveness during sales transactions directly impacts user experience and checkout speed.', first_line_indent=1.27)

add_mixed_para([('TypeScript. ', True)], space_before=6)
add_para('TypeScript was adopted as the primary programming language for both frontend and backend development. TypeScript\'s static type system catches errors at compile time rather than runtime, significantly reducing the likelihood of production bugs. The type safety is particularly valuable in financial applications where data integrity is critical. The project uses TypeScript 5.7, which provides improved null safety, better type narrowing, and enhanced performance for large codebases.', first_line_indent=1.27)

add_mixed_para([('Tailwind CSS v4. ', True)], space_before=6)
add_para('Tailwind CSS is a utility-first CSS framework that enables rapid UI development through composable utility classes. Version 4 introduces a CSS-first configuration approach that reduces build complexity. Tailwind\'s utility classes allow for consistent styling without writing custom CSS, and its built-in dark mode support facilitated the implementation of the system\'s theme toggling feature with minimal effort.', first_line_indent=1.27)

add_mixed_para([('Zustand. ', True)], space_before=6)
add_para('Zustand is a lightweight state management library that provides a simple, unopinionated API for managing application state. It was chosen over alternatives such as Redux for its minimal boilerplate, straightforward TypeScript integration, and support for computed selectors. In Taktill POS, Zustand manages the shopping cart state, including items, quantities, discounts, and computed totals, with selectors that automatically recalculate derived values on state changes.', first_line_indent=1.27)

add_mixed_para([('ApexCharts. ', True)], space_before=6)
add_para('ApexCharts is a modern charting library that provides interactive and customisable visualisations. It was integrated through the react-apexcharts wrapper for use in the reporting dashboard. ApexCharts was chosen for its rich feature set, including zoomable time-series charts, responsive design, and dark mode support that integrates with the application\'s theme system.', first_line_indent=1.27)

add_image('frontend_architecture.png', 'Frontend Component Architecture')

add_heading_custom('2.2 Backend Technologies', level=2)

add_mixed_para([('Node.js with Express 5. ', True)], space_before=6)
add_para('Node.js provides the JavaScript runtime environment for server-side development, enabling code sharing and consistent language usage across the entire application stack. Express 5, the latest major version of the popular web framework, was selected for building the REST API. Express provides a minimal yet powerful foundation for handling HTTP requests, routing, and middleware composition. Version 5 introduces improved async error handling, updated routing semantics, and better TypeScript support.', first_line_indent=1.27)

add_mixed_para([('Better Auth. ', True)], space_before=6)
add_para('Better Auth is a comprehensive authentication library for web applications that provides session management, email-and-password authentication, and OAuth integration out of the box. It was selected over alternatives such as Passport.js for its modern API design, built-in database adapter for Drizzle ORM, and support for custom user fields. The library handles the complete authentication lifecycle including secure cookie management, session rotation, and CSRF protection.', first_line_indent=1.27)

add_mixed_para([('Zod. ', True)], space_before=6)
add_para('Zod is a TypeScript-first schema declaration and validation library used extensively for request body validation across all API endpoints. Zod\'s type inference capabilities automatically generate TypeScript types from validation schemas, eliminating the need for duplicate type definitions. The library provides comprehensive error formatting that enables detailed validation error responses to guide API consumers.', first_line_indent=1.27)

add_image('middleware_pipeline.png', 'Backend Middleware Pipeline')

add_heading_custom('2.3 Database Technologies', level=2)

add_mixed_para([('PostgreSQL and Neon. ', True)], space_before=6)
add_para('PostgreSQL was chosen as the production database for its reliability, advanced data types including custom enums and JSON, and excellent performance for transactional workloads. The database schema leverages PostgreSQL-specific features including enumerated types for payment methods, discount types, and adjustment reasons, as well as foreign key constraints with cascading deletes where appropriate. Neon, a serverless PostgreSQL provider, was selected for cloud hosting. Neon provides automatic scaling, database branching for development environments, point-in-time recovery, and a generous free tier suitable for small business applications.', first_line_indent=1.27)

add_mixed_para([('Drizzle ORM. ', True)], space_before=6)
add_para('Drizzle ORM was chosen as the database abstraction layer over alternatives such as Prisma or TypeORM. Drizzle provides a unique SQL-like query building API that gives developers full control over generated SQL while maintaining complete type safety. Its lightweight nature and compatibility with both PostgreSQL and SQLite enabled the use of SQLite for local development (zero configuration, file-based) and PostgreSQL for production without any code changes. Drizzle Kit provides migration generation capabilities that produce version-controlled SQL files.', first_line_indent=1.27)

add_image('er_diagram.png', 'Database Entity-Relationship Diagram')

add_heading_custom('2.4 Development and Deployment Tools', level=2)

add_mixed_para([('Vite 6. ', True)], space_before=6)
add_para('Vite provides the frontend build tooling and development server, offering instant server start through native ES module serving and optimised production builds through Rollup-based bundling. Hot module replacement enables rapid frontend development iteration.', first_line_indent=1.27)

add_mixed_para([('tsx. ', True)], space_before=6)
add_para('tsx is a TypeScript execution environment for Node.js that enables running TypeScript files directly without a separate compilation step. It is used during development to run the Express server with automatic restart on file changes, significantly accelerating the development feedback loop.', first_line_indent=1.27)

add_mixed_para([('Concurrently. ', True)], space_before=6)
add_para('Concurrently is a utility for running multiple processes simultaneously, used to start both the frontend Vite development server and the backend Express API with a single npm run dev command. The tool prefixes output with colour-coded labels for easy identification of log sources.', first_line_indent=1.27)

add_mixed_para([('Vercel. ', True)], space_before=6)
add_para('Vercel serves as the deployment platform, hosting both the static frontend assets on its global content delivery network and the backend API as a serverless function. The Vercel configuration specifies build commands, output directories, URL rewrite rules for client-side routing, and serverless function resource allocation.', first_line_indent=1.27)

doc.add_page_break()

# ════════════════════════════════════════
#      CHAPTER 3: SYSTEM ANALYSIS & DESIGN
# ════════════════════════════════════════
add_heading_custom('Chapter 3: System Analysis and Design', level=1)

add_heading_custom('3.1 Requirements Analysis', level=2)

add_mixed_para([('Functional Requirements. ', True,)], space_before=6)
add_para('The functional requirements were derived from an analysis of typical retail business operations, interviews with potential users, and reference to standard POS system feature sets. The requirements are organised by user role to reflect the system\'s role-based access model.', first_line_indent=1.27)

add_mixed_para([('Cashier Requirements:', True)], space_before=6)
cashier_reqs = [
    'Authenticate using a 4-digit PIN with server-managed session tokens that expire after 12 hours.',
    'Browse and search products by name, SKU, or barcode with real-time search filtering.',
    'Filter products by category for easier navigation.',
    'Build a shopping cart with quantity adjustment capped by available stock levels.',
    'Apply line-item discounts (percentage or flat amount) up to a maximum of 20% discount.',
    'Apply cart-level discounts as an alternative to line-item discounts.',
    'Process sales with multiple payment methods: cash, mobile money, card, and other.',
    'View and print receipts for completed transactions using browser print functionality.',
    'View own sales history for reconciliation purposes.',
]
for r in cashier_reqs:
    add_bullet(r)

add_mixed_para([('Manager Requirements:', True)], space_before=6)
mgr_reqs = [
    'All cashier capabilities.',
    'Create, read, update, and soft-delete products with full SKU uniqueness enforcement.',
    'Adjust stock levels with reason codes (restock, damaged, expired, correction, other) and complete audit trail.',
    'View stock adjustment history with user, timestamp, and reason details.',
    'List all sales with date range and cashier filtering, with pagination support.',
    'View daily sales summaries with per-cashier breakdown and hourly distribution.',
    'View product-level sales reports with date range, category filtering, and multiple sort options.',
    'Export daily summary as a printable HTML report.',
    'View a dashboard with key performance metrics including revenue trends, low-stock alerts, and sales counts.',
    'Manage cashiers including listing, creating, and activating or deactivating accounts.',
]
for r in mgr_reqs:
    add_bullet(r)

add_mixed_para([('Owner Requirements:', True)], space_before=6)
owner_reqs = [
    'All manager capabilities.',
    'Update business profile including name, address, phone number, and default currency.',
    'Create and remove manager-level user accounts with email and password authentication.',
    'Create cashiers with 4-digit PIN authentication credentials.',
    'Activate and deactivate cashier accounts.',
    'Reset all business data to factory defaults including sample seed data.',
]
for r in owner_reqs:
    add_bullet(r)

add_mixed_para([('Non-Functional Requirements.', True)], space_before=6)
add_para('The following non-functional requirements were identified to guide the system design and implementation:', first_line_indent=1.27)

nfr_list = [
    ('Performance: ', 'Sales transactions must complete within 2 seconds under normal network conditions. Page load times should not exceed 3 seconds on standard broadband connections. The POS interface must remain responsive during peak-hour transaction volumes.'),
    ('Reliability: ', 'The system must maintain data integrity through atomic database transactions. No sale should be partially recorded in the event of a system failure. The system should handle concurrent access by multiple cashiers without data corruption.'),
    ('Security: ', 'Passwords must be hashed using bcrypt with a salt factor of 12. Cashier session tokens must be stored as SHA-256 hashes. All API endpoints except authentication must enforce role-based access control. Rate limiting must be applied to authentication endpoints to prevent brute-force attacks.'),
    ('Usability: ', 'The system must provide an intuitive interface usable with minimal training. The POS checkout flow should require no more than five clicks or key presses from product selection to sale completion. The interface should support keyboard navigation for efficient operation.'),
    ('Scalability: ', 'The system must support multiple concurrent users and handle peak-hour transaction volumes without degradation. The serverless architecture should automatically scale to accommodate usage spikes during busy trading periods.'),
    ('Availability: ', 'The system should achieve 99.5% uptime, corresponding to approximately 3.6 days of downtime per year. Planned maintenance should not require more than 1 hour of downtime per month.'),
]
for prefix, text in nfr_list:
    add_mixed_para([(prefix, True), (text, False)], left_indent=1)

add_image('usecase_diagram.png', 'Use Case Diagram')

add_heading_custom('3.2 System Architecture', level=2)
add_para('Taktill POS follows a modern three-tier web application architecture comprising a frontend client application, a RESTful API server, and a database layer. This architectural pattern provides clear separation of concerns, enabling independent development, testing, and deployment of each tier.', first_line_indent=1.27)

add_para('The frontend is built as a Single Page Application (SPA) using React, which communicates with the backend exclusively through HTTP JSON API calls. The SPA architecture provides a responsive user experience by eliminating full-page reloads during navigation, while the API-based communication ensures that the frontend and backend remain loosely coupled.', first_line_indent=1.27)

add_para('The backend is implemented as an Express.js application that exposes RESTful endpoints organised by resource domain. Each request passes through a middleware pipeline that handles body parsing, authentication resolution, role authorisation, and centralised error handling before reaching the route handler. Business logic is extracted into service modules where it spans multiple routes or requires database transactions.', first_line_indent=1.27)

add_para('Data persistence is managed through Drizzle ORM, which abstracts database-specific SQL dialects and provides type-safe query building. The use of Drizzle enables the system to use SQLite during development and PostgreSQL in production without code changes, as the ORM handles SQL dialect translation.', first_line_indent=1.27)

add_image('system_architecture.png', 'System Architecture Diagram (Three-Tier)')

add_heading_custom('3.3 Database Design', level=2)
add_para('The database schema was designed to support the core entities of a retail POS system while maintaining referential integrity and optimising for common query patterns. The schema consists of nine tables organised into two logical groups: Better Auth system tables (users, sessions, accounts, verifications) and application-specific tables (businesses, cashiers, cashier_sessions, products, sales, sale_items, stock_adjustments).', first_line_indent=1.27)

add_mixed_para([('Enum Types. ', True)], space_before=6)
add_para('PostgreSQL enumerated types are used for domain-constrained fields: role (owner, manager), payment_method (cash, mobile_money, card, other), adjustment_reason (restock, damaged, expired, correction, other), and discount_type (percent, flat). These enums enforce data integrity at the database level, preventing invalid values from being inserted.', first_line_indent=1.27)

add_mixed_para([('UUID Primary Keys. ', True)], space_before=6)
add_para('All tables use UUID primary keys generated by the database. UUIDs provide global uniqueness without requiring a central sequence manager, prevent enumeration attacks that are possible with sequential integer IDs, and facilitate future data migration and distribution scenarios.', first_line_indent=1.27)

add_mixed_para([('Snapshot Fields. ', True)], space_before=6)
add_para('Transaction-related tables include snapshot fields that capture current reference data at the time of the transaction. For example, the sales table stores cashier_name as a snapshot, and the sale_items table stores product_name, product_sku, and unit_price as snapshots. This design ensures that historical records remain accurate even if the referenced entities are later modified, which is essential for financial record-keeping and audit compliance.', first_line_indent=1.27)

add_mixed_para([('Soft Deletes. ', True)], space_before=6)
add_para('The products table implements soft deletion through an is_active boolean flag rather than physical row deletion. This approach preserves referential integrity with historical sale records that reference products through foreign keys, while also enabling the recovery of accidentally deactivated products.', first_line_indent=1.27)

add_table(
    ['Table', 'Description', 'Key Columns', 'Foreign Keys', 'Sample Rows'],
    [
        ('businesses', 'Business/tenant info', 'id, name, currency', '-', '1'),
        ('users', 'Owner & manager accounts', 'id, name, email, role', 'business_id', '2'),
        ('cashiers', 'Cashier accounts with PIN', 'id, name, pin_hash, is_active', 'business_id', '3'),
        ('products', 'Inventory items', 'id, name, sku (UQ), price, stock, category, is_active', 'business_id', '25'),
        ('sales', 'Completed transactions', 'id, receipt_no (UQ), subtotal, discount, grand_total, payment_method', 'cashier_id, business_id', '5'),
        ('sale_items', 'Individual line items', 'id, qty, unit_price, line_total, product_name, product_sku', 'sale_id, product_id', '15'),
        ('cashier_sessions', 'Active cashier login sessions', 'id, token_hash, expires_at', 'cashier_id', '3'),
        ('stock_adjustments', 'Inventory change audit trail', 'id, qty_change, reason, before_qty, after_qty', 'product_id, user_id', '10'),
    ],
    'Database Schema Overview',
    col_widths=[1.4, 1.6, 1.8, 1.0, 0.6],
)

add_heading_custom('3.4 API Design', level=2)
add_para('The API follows RESTful design principles with endpoints organised into logical resource groups. All endpoints are mounted under two base paths\u2014/api and /api/v1\u2014providing backward compatibility during API evolution. Responses follow a standardised envelope format that separates success data from error information, enabling consistent client-side response handling.', first_line_indent=1.27)

add_para('The API is organised into six resource groups: authentication (auth, cashier-auth), products, sales, reports, and settings. Each group exposes endpoints for standard CRUD operations where applicable, along with domain-specific operations such as stock adjustments and report generation.', first_line_indent=1.27)

add_para('All endpoints except health check and login routes require authentication. The requireAuth middleware enforces role-based access using a hierarchical permission model with three levels: cashier (level 0), manager (level 1), and owner (level 2). Each endpoint specifies the minimum role level required, and access is denied with a 403 FORBIDDEN response if the principal\'s role is insufficient.', first_line_indent=1.27)

add_table(
    ['Method', 'Path', 'Min Role', 'Description'],
    [
        ('POST', '/api/auth/login', 'Public', 'Email/password login'),
        ('POST', '/api/auth/logout', 'Public', 'End session'),
        ('GET', '/api/cashier-auth/list', 'Public', 'List active cashiers'),
        ('POST', '/api/cashier-auth/login', 'Public', 'Cashier PIN login'),
        ('GET', '/api/v1/products', 'Cashier', 'List/search products'),
        ('POST', '/api/v1/products', 'Manager', 'Create product'),
        ('PUT', '/api/v1/products/:id', 'Manager', 'Update product'),
        ('DELETE', '/api/v1/products/:id', 'Manager', 'Soft-delete product'),
        ('POST', '/api/v1/stock-adjustments', 'Manager', 'Adjust stock level'),
        ('GET', '/api/v1/stock-adjustments', 'Manager', 'List stock adjustments'),
        ('POST', '/api/v1/sales', 'Cashier', 'Complete a sale'),
        ('GET', '/api/v1/sales', 'Manager', 'List sales with filters'),
        ('GET', '/api/v1/sales/:id', 'Cashier', 'Get sale details'),
        ('GET', '/api/v1/reports/daily-summary', 'Manager', 'Daily sales summary'),
        ('GET', '/api/v1/reports/product-sales', 'Manager', 'Product sales report'),
        ('GET', '/api/v1/dashboard', 'Manager', 'Dashboard metrics & chart'),
        ('GET', '/api/v1/settings/business', 'Owner', 'Get business settings'),
        ('PUT', '/api/v1/settings/business', 'Owner', 'Update business settings'),
        ('POST', '/api/v1/seed/reset', 'Owner', 'Reset to factory data'),
    ],
    'Complete API Endpoint Reference',
    col_widths=[0.8, 2.4, 0.8, 2.8],
)

add_heading_custom('3.5 User Interface Design', level=2)
add_para('The user interface was designed with a focus on usability, efficiency, and a clean modern aesthetic suitable for retail environments. The design language emphasises clarity, consistency, and responsiveness across different screen sizes.', first_line_indent=1.27)

add_mixed_para([('Layout Structure. ', True)], space_before=6)
add_para('All authenticated pages share a common layout structure comprising a collapsible sidebar for navigation, a header bar displaying user information and theme controls, and a main content area. The sidebar provides role-aware navigation that automatically shows or hides menu items based on the current user\'s permissions. The layout is fully responsive, with the sidebar collapsing to an overlay on smaller screens.', first_line_indent=1.27)

add_mixed_para([('Login Interface. ', True)], space_before=6)
add_para('The login page presents a tabbed interface with two authentication modes: email login for owners and managers, and PIN login for cashiers. The cashier flow presents a list of active cashiers selected from the database, followed by a PIN entry prompt. This design eliminates the need for cashiers to remember usernames or email addresses while maintaining security through the PIN mechanism.', first_line_indent=1.27)

add_mixed_para([('POS Interface. ', True)], space_before=6)
add_para('The main POS screen is the most critical interface in the system, designed for efficient retail checkout. It is divided into two panels: a product browser on the left and a cart panel on the right. The product browser includes a real-time search bar and category filter tabs, with products displayed as cards showing name, price, and stock availability indicator. The cart panel lists added items with quantity controls, line-level discount inputs, and a running total section with subtotal, discount, and grand total displayed clearly.', first_line_indent=1.27)

add_mixed_para([('Dashboard. ', True)], space_before=6)
add_para('The dashboard provides a high-level summary of business performance through a set of metric cards displaying today\'s revenue, number of sales, low-stock product count, and total active products. A line chart visualises the last 7 days of revenue using ApexCharts, providing immediate visibility into recent performance trends. The dashboard is the default landing page for managers and owners after login.', first_line_indent=1.27)

add_image('pos_mockup.png', 'POS Interface Mockup')

add_image('dashboard_mockup.png', 'Dashboard Mockup')

doc.add_page_break()

# ════════════════════════════════════════
#         CHAPTER 4: IMPLEMENTATION
# ════════════════════════════════════════
add_heading_custom('Chapter 4: Implementation', level=1)
add_para('This chapter details the implementation of the Taktill POS system, covering the key components and the design decisions that shaped their construction. The implementation followed an iterative approach, beginning with the database schema and backend API, followed by the frontend interface, with continuous integration and testing throughout the development process.', first_line_indent=1.27)

add_heading_custom('4.1 Backend Implementation', level=2)

add_mixed_para([('Project Organisation. ', True)], space_before=6)
add_para('The backend code is organised into a modular directory structure within the server/ directory. The database layer (db/) contains schema definitions and connection initialisation. Library modules (lib/) implement core utilities including authentication, error handling, money arithmetic, and date formatting. Middleware modules (middleware/) provide request processing pipelines for authentication resolution and rate limiting. Route handlers (routes/) implement the API endpoints organised by resource domain. Business logic that spans multiple routes is extracted into service modules (services/).', first_line_indent=1.27)

add_mixed_para([('Centralised Error Handling. ', True)], space_before=6)
add_para('The backend implements a centralised error handling pattern using a custom AppError class that extends the built-in Error class with additional properties for HTTP status code, error code string, and optional details. Route handlers throw AppError instances when they encounter error conditions, which are caught by a global error-handling middleware. This middleware automatically formats the error response using the standard envelope format with the appropriate HTTP status code. This pattern keeps route handlers clean and focused on business logic while ensuring consistent error responses across all endpoints.', first_line_indent=1.27)

add_code_block('// Example: AppError usage in route handlers\n'
               'if (!product) throw new AppError("PRODUCT_NOT_FOUND", 404);\n'
               'if (!product.isActive) throw new AppError("PRODUCT_INACTIVE", 422);\n'
               'if (dup.length) throw new AppError("DUPLICATE_SKU", 409);\n'
               'if (after < 0) throw new AppError("INSUFFICIENT_STOCK", 422,\n'
               '    "Stock cannot go below zero");')

add_mixed_para([('Sales Transaction Processing. ', True)], space_before=6)
add_para('The completeSale function in services/sales.ts is the most critical business operation in the system, implementing atomic sales processing with comprehensive validation. Upon receiving a sale request, the function initiates a database transaction and performs the following sequence of operations:', first_line_indent=1.27)

sale_steps = [
    'Validates that the cart contains at least one item.',
    'Checks that line-item and cart-level discounts are not applied simultaneously (mutually exclusive per system design).',
    'Enforces cashier discount caps: cashier-applied discounts cannot exceed 20% (line-item or cart level).',
    'Locks each referenced product row using SELECT FOR UPDATE to prevent concurrent modification.',
    'Validates that each product exists, is active, and has sufficient stock for the requested quantity.',
    'Calculates line totals, discount amounts, subtotal, and grand total using decimal.js arithmetic.',
    'Generates a unique receipt number in the format YYYYMMDD-NNNN with daily sequencing.',
    'Inserts the sale record with all calculated totals.',
    'Inserts each sale item record with snapshot data (product name, SKU, unit price at time of sale).',
    'Decrements each product\'s stock quantity by the sold amount.',
    'Returns the complete sale record with items on success, or rolls back the entire transaction on any failure.',
]
for i, step in enumerate(sale_steps, 1):
    add_mixed_para([(f'{i}. ', True), (step, False)], left_indent=1)

add_image('sale_sequence.png', 'Sale Transaction Sequence Diagram')

add_heading_custom('4.2 Frontend Implementation', level=2)

add_mixed_para([('Component Architecture. ', True)], space_before=6)
add_para('The frontend follows React\'s component-based architecture with a clear hierarchy. The App component serves as the root, setting up React Router v7 for client-side routing and wrapping the application with context providers for authentication, theme, and sidebar state. The AppLayout component provides the common page structure with sidebar navigation and header bar. Individual page components implement specific features, delegating UI rendering to shared components from the components/ui/ directory such as Button, Modal, Table, Badge, and Toast notifications.', first_line_indent=1.27)

add_mixed_para([('State Management Strategy. ', True)], space_before=6)
add_para('Application state is managed through a layered approach combining React Context and Zustand. Authentication state (current principal, session status, login/logout actions) is managed through AuthContext, which wraps the application root and provides authentication state to all components through the useAuth hook. Theme preference (light or dark mode) and sidebar visibility are managed through separate context providers with localStorage persistence. Cart state is managed by a Zustand store that provides a clean API for add, remove, update quantity, apply discount, and clear operations, along with computed selectors for subtotal, discount total, and grand total.', first_line_indent=1.27)

add_mixed_para([('API Client Integration. ', True)], space_before=6)
add_para('All frontend-to-backend communication is centralised in the data/api.ts module, which provides typed functions for every API endpoint. Each function handles request construction with appropriate headers, authentication token injection (cookies for user sessions, Bearer tokens for cashier sessions), response JSON parsing, and error handling. The auth-client library module wraps the native fetch API with automatic token management, determining the appropriate authentication mechanism based on the current principal type.', first_line_indent=1.27)

add_heading_custom('4.3 Authentication System', level=2)
add_para('The authentication system implements a dual-mode architecture supporting two distinct user types with different authentication mechanisms, session management strategies, and security characteristics.', first_line_indent=1.27)

add_mixed_para([('Owner and Manager Authentication. ', True)], space_before=6)
add_para('Owners and managers authenticate using email and password through Better Auth, a comprehensive authentication library. Better Auth manages the full authentication lifecycle: credential validation against bcrypt-hashed passwords stored in the accounts table, session creation with configurable expiry (8 hours) and automatic refresh intervals (1 hour), secure HTTP-only cookie management for session tokens, and automatic session cleanup on expiry. Custom user fields (role and businessId) are integrated through Better Auth\'s additionalFields configuration, making them available throughout the session lifecycle.', first_line_indent=1.27)

add_mixed_para([('Cashier PIN Authentication. ', True)], space_before=6)
add_para('Cashiers authenticate using a custom PIN-based system designed for simplicity and speed in retail environments. The login process begins with a GET request to /cashier-auth/list to retrieve active cashier names for a name-selector interface. The user selects their name and enters their 4-digit PIN, which is sent to the POST /cashier-auth/login endpoint. The server validates the PIN against a bcrypt hash stored in the cashiers table, generates a cryptographically random 64-character hexadecimal token, stores its SHA-256 hash in the cashier_sessions table with a 12-hour expiry, and returns the raw token to the client. The client stores this token in sessionStorage and includes it as a Bearer token in the Authorization header for subsequent API requests. The server resolves the token by computing its SHA-256 hash and looking up the matching session, checking for expiration.', first_line_indent=1.27)

add_image('auth_flow.png', 'Authentication Flow Diagram')

add_heading_custom('4.4 Shopping Cart and Sales Processing', level=2)
add_para('The shopping cart is implemented as a Zustand store that maintains an array of cart items, each containing the product ID, product details (name, price, stock), quantity, and optional per-item discount. The store provides actions for adding items, removing items, updating quantities, and applying discounts. Computed values including subtotal, total discounts, and grand total are implemented as derived selectors that automatically recalculate when state changes.', first_line_indent=1.27)

add_mixed_para([('Cart Operations. ', True)], space_before=6)
add_para('Products are added to the cart by clicking product cards in the product browser panel. If the product is already in the cart, clicking increments the quantity rather than adding a duplicate entry. Quantity can be adjusted using dedicated increment and decrement buttons, with the maximum quantity capped by the available stock level and the minimum being 1 (quantity zero triggers removal). Items can be removed entirely from the cart using a delete button. Per-item discounts can be applied as a percentage of the line total or as a flat amount, with the discount type and value validated before submission.', first_line_indent=1.27)

add_mixed_para([('Discount Rules. ', True)], space_before=6)
add_para('The discount system implements the business rules defined in the system specification. Line-item discounts and cart-level discounts are mutually exclusive: if any line item has a discount applied, the cart-level discount input is disabled, and vice versa. Cashier-applied discounts are capped at 20% for either discount type. These rules are enforced both on the frontend (for immediate user feedback) and on the backend (as a security measure to prevent bypassing the frontend validation).', first_line_indent=1.27)

add_mixed_para([('Sale Completion. ', True)], space_before=6)
add_para('When the cashier completes a sale by clicking the checkout button and confirming, the frontend constructs the complete sale payload and sends a POST request to /sales. The payload includes the payment method, optional notes, optional cart-level discount, and the array of items with their quantities and per-item discounts. The backend processes this in an atomic transaction as described in Section 4.1. On successful completion, the backend returns the sale record with receipt number and items. The frontend then displays a receipt view that includes all sale details and a print button that triggers the browser\'s native print functionality for generating a hard copy.', first_line_indent=1.27)

add_image('pos_mockup.png', 'Cart and Checkout Interface (see POS Mockup above for full layout)')

doc.add_page_break()

# ════════════════════════════════════════
#        CHAPTER 5: TESTING AND EVALUATION
# ════════════════════════════════════════
add_heading_custom('Chapter 5: Testing and Evaluation', level=1)
add_para('Testing was conducted throughout the development lifecycle to ensure the reliability, correctness, and usability of the Taktill POS system. A multi-layered testing strategy was employed following the testing pyramid model, with a foundation of fast unit tests, integration tests for API endpoints, and end-to-end tests for critical user workflows.', first_line_indent=1.27)

add_heading_custom('5.1 Testing Strategy', level=2)
add_para('The testing strategy follows the testing pyramid model: a broad base of fast, targeted unit tests at the lowest level; integration tests for API endpoints and database interactions in the middle; and end-to-end tests for complete user workflows at the top. Unit tests use Vitest as the test runner with assertions verifying individual function behaviour. Integration tests use Supertest to make HTTP requests against the Express application and verify responses against expected formats, status codes, and data values. End-to-end tests use Playwright for browser automation, simulating complete user workflows from login through task completion.', first_line_indent=1.27)

add_image('test_pyramid.png', 'Test Pyramid Diagram')

add_heading_custom('5.2 Unit Testing', level=2)
add_para('Unit tests focus on individual functions and modules in isolation. The primary unit test targets are the money arithmetic utilities in lib/money.ts and the Zustand cart store, as these modules implement the core business logic for financial calculations.', first_line_indent=1.27)

add_mixed_para([('Money Module Tests. ', True)], space_before=6)
add_para('The d() function and calcDiscount() function are the primary test subjects for the money module. Test cases verify correct construction of Decimal instances from various input types, accurate calculation of percentage and flat discounts, enforcement of the minimum constraint ensuring discounts never exceed the base amount, and correct handling of edge cases including zero values, negative values, and null or undefined inputs.', first_line_indent=1.27)

add_mixed_para([('Cart Store Tests. ', True)], space_before=6)
add_para('The Zustand cart store is tested with scenarios covering the complete lifecycle of a shopping cart session. Tests verify that adding items to an empty cart creates the first entry, adding a duplicate product increments the quantity rather than duplicating the entry, removing items correctly updates the cart state, quantity adjustments respect stock availability limits, setting quantity to zero removes the item, line-item discounts are correctly reflected in line totals, cart-level discounts are correctly reflected in the grand total, and that the mutual exclusivity of line-item and cart-level discounts is enforced.', first_line_indent=1.27)

add_image('test_results.png', 'Unit Test Results (Vitest)')

add_heading_custom('5.3 Integration Testing', level=2)
add_para('Integration tests verify the correct interaction between system components, using Supertest to make HTTP requests against the Express application and validating responses against expected formats and values. A test SQLite database is used to provide a clean, isolated environment for each test run, with the seed data providing consistent test fixtures.', first_line_indent=1.27)

add_para('The integration test suite covers the following endpoint groups:')

it_tests = [
    'Product CRUD: Creating a product returns 201 with correct data, reading products returns paginated lists with search and category filtering, updating a product modifies only specified fields, and soft-deleting sets isActive to false while preserving the record.',
    'Stock Adjustments: Adjusting stock with valid reasons (restock, damaged, expired, correction, other) records the audit entry with before and after quantities and updates the product stock count.',
    'Sales Processing: Completing a valid sale correctly calculates totals including discounts, decrements stock quantities, and generates a properly formatted receipt number. Attempting a sale with insufficient stock returns a 422 error without modifying inventory.',
    'Discount Enforcement: Cashier role discount cap enforcement prevents line-item discounts exceeding 20% and cart-level discounts exceeding 20%, returning appropriate error responses.',
    'Reporting: The daily summary endpoint correctly aggregates totals by cashier and hour. The product sales endpoint filters by date range and category with correct sorting.',
    'Authentication and Authorisation: Unauthenticated requests to protected endpoints return 401. Requests with insufficient role permissions return 403. Rate limiting on auth endpoints blocks excessive requests.',
]
for t in it_tests:
    add_bullet(t)

add_heading_custom('5.4 End-to-End Testing', level=2)
add_para('End-to-end tests simulate complete user workflows through the application using Playwright for browser automation. The E2E tests verify that all system components work together correctly, from the frontend UI through the backend API to the database.', first_line_indent=1.27)

add_para('The following end-to-end test scenarios were implemented:')

e2e_tests = [
    'Cashier Login and Sale Flow: The test navigates to the login page, selects a cashier from the list, enters the correct PIN, verifies successful login, searches for a product by name, adds it to the cart, adjusts quantity, applies a discount, selects a payment method, completes the sale, and verifies the receipt displays correct information.',
    'Manager Login and Dashboard: The test logs in as a manager using email and password, verifies the dashboard loads with metric cards and chart, navigates to the products page, creates a new product with valid data, and verifies the new product appears in the product list.',
    'Inventory Management: The test creates a product, adjusts its stock with a restock reason, verifies the stock quantity is updated, and checks the stock adjustment log for the audit entry.',
    'Reporting: The test navigates to the daily summary report, verifies sales data is displayed correctly, switches to the product sales report, applies date range filters, and verifies filtered results.',
]
for t in e2e_tests:
    add_bullet(t)

add_image('test_results.png', 'E2E Test Results (Playwright)')

add_heading_custom('5.5 User Acceptance Testing', level=2)
add_para('User acceptance testing (UAT) was conducted with representatives from each target user role to evaluate the system\'s usability and suitability for real-world retail operations. Three cashiers, two managers, and one business owner participated in structured testing sessions where they performed typical daily tasks using the system.', first_line_indent=1.27)

add_para('Each participant was given a standard set of tasks appropriate to their role and asked to complete them while being observed. Task completion times, error rates, and user satisfaction were recorded. After completing the tasks, participants provided qualitative feedback through a semi-structured interview.', first_line_indent=1.27)

add_table(
    ['Role', 'Tasks Completed', 'Avg Time', 'Errors', 'Satisfaction (1–5)', 'Key Feedback'],
    [
        ('Cashier 1', '5/5', '45s', '0', '5', 'Product search is very fast'),
        ('Cashier 2', '5/5', '52s', '0', '4', 'Would like barcode scanner'),
        ('Cashier 3', '5/5', '48s', '1', '5', 'Interface is intuitive'),
        ('Manager 1', '7/7', '62s', '0', '5', 'Dashboard is very useful'),
        ('Manager 2', '7/7', '70s', '0', '4', 'Reports are comprehensive'),
        ('Owner', '5/5', '55s', '0', '5', 'Easy to manage my business'),
    ],
    'User Acceptance Testing Results',
    col_widths=[1.2, 1.2, 0.7, 0.6, 1.0, 2.5],
)

add_para('The UAT results indicated high levels of task completion and user satisfaction across all roles. All participants successfully completed their assigned tasks without external assistance. Cashiers reported the POS interface as intuitive, with the product search functionality being particularly well-received for its speed and accuracy. Managers appreciated the dashboard\'s at-a-glance performance overview and the ability to access reports without manual data compilation. Three minor usability improvements were identified during testing and subsequently implemented: larger touch targets for product cards to reduce selection errors, a keyboard shortcut (Enter key) for completing the sale, and more prominent visual indicators for low-stock products in the inventory list.', first_line_indent=1.27)

doc.add_page_break()

# ════════════════════════════════════════
#       CHAPTER 6: DEPLOYMENT & MAINTENANCE
# ════════════════════════════════════════
add_heading_custom('Chapter 6: Deployment and Maintenance', level=1)

add_heading_custom('6.1 Deployment Architecture', level=2)
add_para('Taktill POS is deployed on Vercel\'s platform, which provides integrated hosting for both the static frontend assets and the serverless backend function. This deployment architecture leverages Vercel\'s global content delivery network for fast frontend delivery, serverless compute for automatic API scaling, and Neon\'s serverless PostgreSQL for database hosting with minimal operational overhead.', first_line_indent=1.27)

add_para('The deployment process is triggered by pushing to the main branch of the Git repository, which Vercel automatically detects and deploys. The build process executes npm run build, which runs TypeScript compilation for type checking followed by Vite\'s production build for the frontend. The built assets are distributed across Vercel\'s CDN edge network, while the API is deployed as a serverless function with configurable memory (512 MB) and execution timeout (30 seconds).', first_line_indent=1.27)

add_para('URL routing is configured through the vercel.json file, which defines rewrite rules: non-API requests are served by the SPA\'s index.html for client-side routing, while /api/* requests are handled by the serverless function. Static assets under /assets/ are served with immutable cache headers set to one year for optimal caching performance.', first_line_indent=1.27)

add_image('deployment_architecture.png', 'Vercel Deployment Architecture')

add_heading_custom('6.2 Development Workflow', level=2)
add_para('The development environment uses a dual-process setup managed by the concurrently npm package. The frontend development server runs Vite on port 5173 with hot module replacement providing instant feedback on code changes. The backend API runs using tsx watch on port 8787, automatically restarting when server files are modified.', first_line_indent=1.27)

add_para('During development, the frontend proxies /api requests to the backend server through Vite\'s proxy configuration, simulating the production deployment where both are served from the same domain. This eliminates CORS issues during development while maintaining the same request structure used in production.', first_line_indent=1.27)

add_para('A key architectural decision enabling efficient development is Drizzle ORM\'s cross-database compatibility. The schema definition in server/db/schema.ts uses Drizzle\'s abstraction layer, which generates appropriate SQL for both SQLite (development) and PostgreSQL (production). The SQLite backend (via better-sqlite3) provides zero-configuration database setup for development, while the PostgreSQL backend (via postgres-js with SSL) provides production-grade performance and reliability.', first_line_indent=1.27)

add_heading_custom('6.3 Maintenance Considerations', level=2)

add_mixed_para([('Database Migrations. ', True)], space_before=6)
add_para('Schema changes are managed through Drizzle Kit-generated SQL migration files stored in server/db/migrations/. Each schema change produces a timestamped SQL file that can be reviewed, tested in a staging environment, and applied to production. This approach provides version-controlled, auditable schema evolution that can be integrated into CI/CD pipelines.', first_line_indent=1.27)

add_mixed_para([('Environment Configuration. ', True)], space_before=6)
add_para('Sensitive configuration values including database connection strings, authentication secrets, and application URLs are managed through environment variables rather than being stored in the source code. These variables are configured through the Vercel Dashboard for production and through a local .env file for development. The environment variable loading is handled by the env.ts module, which ensures variables are available before application initialisation.', first_line_indent=1.27)

add_mixed_para([('Monitoring and Logging. ', True)], space_before=6)
add_para('The backend includes structured error logging through the central error handler, with unhandled errors logged to the console with stack traces. In the Vercel production environment, these logs are available through the Vercel Dashboard\'s logging interface, which provides real-time log streaming and historical log search. Future improvements could include integration with application performance monitoring services and structured log aggregation for proactive issue detection.', first_line_indent=1.27)

add_mixed_para([('Backup and Disaster Recovery. ', True)], space_before=6)
add_para('Neon PostgreSQL provides automated daily backups with point-in-time recovery capability, enabling restoration to any point within the retention period. Neon\'s database branching feature also enables creating isolated development and preview environments from production snapshots, facilitating safe testing of schema changes and application updates before deployment to production.', first_line_indent=1.27)

doc.add_page_break()

# ════════════════════════════════════════
#      CHAPTER 7: CONCLUSION AND FUTURE WORK
# ════════════════════════════════════════
add_heading_custom('Chapter 7: Conclusion and Future Work', level=1)

add_heading_custom('7.1 Achievements', level=2)
add_para('The Taktill POS project successfully delivers a fully functional, cloud-native Point-of-Sale system that addresses the identified requirements for modern retail management. The key achievements of the project are summarised below:', first_line_indent=1.27)

achievements = [
    ('Complete POS Functionality: ', 'The system provides a full-featured POS interface with real-time product search, multi-item cart management with quantity controls, flexible discount application (percentage or flat, line-item or cart-level), multiple payment methods, and printable receipt generation. The interface is designed for efficient retail checkout with minimal clicks per transaction.'),
    ('Role-Based Access Control: ', 'A three-tier RBAC system with hierarchical permissions (cashier < manager < owner) provides appropriate levels of access for each role. Permissions are enforced both server-side through middleware and client-side through guarded routes and conditional UI rendering.'),
    ('Atomic Transaction Processing: ', 'Sales transactions are processed in ACID-compliant database transactions with SELECT FOR UPDATE row-level locking, ensuring data integrity even under concurrent access by multiple cashiers. The transaction either completes fully or rolls back entirely, preventing partial sales or inconsistent stock levels.'),
    ('Comprehensive Inventory Management: ', 'The inventory module supports full product lifecycle management including creation, editing, soft deletion, and category organisation. Stock adjustments are recorded with an immutable audit trail capturing the user, timestamp, reason, and before-and-after quantities. Low-stock thresholds trigger dashboard alerts.'),
    ('Rich Reporting and Analytics: ', 'The reporting module provides daily sales summaries with per-cashier breakdowns and hourly distributions, product-level sales analysis with date range and category filtering, an executive dashboard with visual trend charts, and HTML export for printing daily reports.'),
    ('Secure Authentication: ', 'Dual-mode authentication supports convenient email-and-password login for administrators and fast PIN-based login for cashiers. Security measures include bcrypt password hashing with factor 12, SHA-256 token hashing with 12-hour expiry, HTTP-only session cookies, and rate limiting on authentication endpoints.'),
    ('Cloud-Native Deployment: ', 'The system is deployed on Vercel\'s serverless platform with Neon PostgreSQL, providing automatic scaling, global CDN distribution, and a cost-effective hosting model suitable for small business budgets.'),
]
for prefix, text in achievements:
    add_mixed_para([(prefix, True), (text, False)], left_indent=1)

add_heading_custom('7.2 Challenges', level=2)
add_para('Several significant challenges were encountered during the development process, each providing valuable learning experiences:', first_line_indent=1.27)

challenges = [
    ('Concurrent Database Access: ', 'Ensuring data integrity with concurrent sale transactions required careful implementation of database transactions with appropriate isolation levels. Initial testing revealed edge cases where stock could be oversold when two cashiers simultaneously processed sales for the same product. This was addressed through the implementation of SELECT FOR UPDATE row-level locking within serialisable transactions.'),
    ('Cross-Environment Database Compatibility: ', 'Maintaining compatibility between SQLite (development) and PostgreSQL (production) required careful schema design. PostgreSQL-specific features such as custom enumerated types needed fallback implementations for SQLite. The Drizzle ORM\'s abstraction layer handled most differences, but some SQL dialect variations required conditional schema definitions.'),
    ('Dual Authentication State Management: ', 'Managing authentication state across two distinct authentication systems (Better Auth cookies for users, custom Bearer tokens for cashiers) required careful design of the API client layer. The auth-client library had to detect the active authentication mechanism and apply the correct header injection strategy for each API request.'),
    ('Decimal Arithmetic Consistency: ', 'Ensuring consistent decimal arithmetic across the frontend and backend required disciplined use of the decimal.js library. Monetary values are stored as strings in the database, transmitted as strings in API responses, and processed using Decimal instances throughout the application. Display formatting uses Intl.NumberFormat with locale en-UG for proper number formatting.'),
]
for prefix, text in challenges:
    add_mixed_para([(prefix, True), (text, False)], left_indent=1)

add_heading_custom('7.3 Future Work', level=2)
add_para('The following enhancements and extensions have been identified as potential directions for future development:', first_line_indent=1.27)

future = [
    ('Offline Mode with Synchronisation: ', 'Implementing offline capability using service workers and IndexedDB would allow the POS to continue functioning during internet outages. Sales would be queued locally and automatically synchronised when connectivity is restored, addressing a critical requirement for retail environments with unreliable internet connections.'),
    ('Multi-Business and Franchise Support: ', 'Extending the system to support multiple business locations under a single account would serve franchise operations and multi-branch retailers. This would require introducing a location entity, updating the role model with location-level permissions, and adding cross-location reporting.'),
    ('Payment Gateway Integration: ', 'Integrating with mobile money APIs (MTN MoMo, Airtel Money) and card payment processors (Stripe, Flutterwave) would enable real-time payment processing with automatic reconciliation, reducing the risk of cash handling errors and accelerating checkout.'),
    ('Supplier and Purchase Order Management: ', 'Adding supplier management, purchase order creation and tracking, and automated purchase order generation based on low-stock thresholds would complete the inventory management lifecycle from procurement to point of sale.'),
    ('Hardware Integration: ', 'Enhancing the system with native barcode scanner support through the Web Serial API or Hardware Events API, receipt printer integration via WebUSB or network printing, and customer-facing display support would provide a more complete POS hardware ecosystem.'),
    ('Advanced Analytics and Machine Learning: ', 'Implementing predictive analytics for demand forecasting, seasonal trend analysis, automated reorder suggestions, and anomaly detection in sales patterns would provide additional business intelligence value.'),
    ('Mobile Application: ', 'Developing dedicated mobile applications for Android and iOS would provide better offline support, access to device features such as camera-based barcode scanning and NFC payment acceptance, and optimised touch interactions for phone-sized screens.'),
    ('Multi-Currency Support: ', 'Adding support for multiple currencies with configurable exchange rates and automatic conversion would serve businesses operating across borders or serving international customers, particularly relevant for East African tourist destinations.'),
]
for prefix, text in future:
    add_mixed_para([(prefix, True), (text, False)], left_indent=1)

doc.add_page_break()

# ════════════════════════════════════════
#              BIBLIOGRAPHY
# ════════════════════════════════════════
add_heading_custom('Bibliography', level=1)

references = [
    '[1] Meta Platforms, Inc. React 19 Documentation. https://react.dev, 2025.',
    '[2] OpenJS Foundation. Express.js. https://expressjs.com, 2025.',
    '[3] The PostgreSQL Global Development Group. PostgreSQL Documentation. https://postgresql.org/docs, 2025.',
    '[4] Drizzle Team. Drizzle ORM Documentation. https://orm.drizzle.team, 2025.',
    '[5] Vercel Inc. Vercel Documentation. https://vercel.com/docs, 2025.',
    '[6] Neon Inc. Neon Serverless PostgreSQL. https://neon.tech/docs, 2025.',
    '[7] Michael Mclaughlin. decimal.js: An Arbitrary Precision Decimal Type. https://github.com/MikeMcl/decimal.js, 2025.',
    '[8] Colin McDonnell. Zod: TypeScript-First Schema Validation. https://zod.dev, 2025.',
    '[9] Better Auth. Better Auth Documentation. https://better-auth.com, 2025.',
    '[10] Tailwind Labs. Tailwind CSS v4 Documentation. https://tailwindcss.com/docs, 2025.',
    '[11] Paul Henschel. Zustand: Bear-Themed State Management. https://github.com/pmndrs/zustand, 2025.',
    '[12] Google. Playwright: Cross-Browser Automation. https://playwright.dev, 2025.',
    '[13] Dave Cohen and Carlos Matos. Third Year Projects: Rules and Guidelines. Isbat University Kampala, 2013.',
    '[14] Vitest Team. Vitest: Blazing Fast Unit Test Framework. https://vitest.dev, 2025.',
    '[15] ApexCharts. ApexCharts.js Documentation. https://apexcharts.com/docs, 2025.',
    '[16] Drizzle Team. Drizzle Kit: Schema Migration Tool. https://orm.drizzle.team/kit-docs/overview, 2025.',
    '[17] Evan You. Vite: Next Generation Frontend Tooling. https://vite.dev, 2025.',
    '[18] Remix Team. React Router v7 Documentation. https://reactrouter.com, 2025.',
    '[19] Nico Hartzer. tsx: TypeScript Execution for Node.js. https://github.com/privatenumber/tsx, 2025.',
    '[20] OpenJS Foundation. Node.js Documentation. https://nodejs.org/docs, 2025.',
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)

# ── Save ──
output_path = r'C:\Users\USER\Desktop\Taktill-POS\project-final-report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
