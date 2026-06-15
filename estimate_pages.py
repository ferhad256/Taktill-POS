import math, os, re
from docx import Document
from lxml import etree

doc = Document(r'C:\Users\USER\Desktop\Taktill-POS\project-final-report.docx')

page_h_in = 9.69
line_h = 14.4 / 72 * 1.5

def emu_to_in(v):
    return (v or 0) / 914400

chars_per_line = 63
current_y = 0
total_pages = 1

for para in doc.paragraphs:
    t = para.text
    sp_b = emu_to_in(para.paragraph_format.space_before)
    sp_a = emu_to_in(para.paragraph_format.space_after)
    has_img = para._element.findall(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    
    if has_img:
        h = 4.0
    elif not t.strip():
        h = line_h * 0.5
    else:
        lines = max(1, math.ceil(len(t) / chars_per_line))
        h = lines * line_h + sp_b + sp_a
    
    current_y += h
    if current_y > page_h_in:
        total_pages += 1
        current_y = h

for tbl in doc.tables:
    th = len(tbl.rows) * 0.2 + 0.3
    current_y += th
    if current_y > page_h_in:
        total_pages += 1
        current_y = th

# Count explicit page breaks
xml = etree.tostring(doc.element, encoding='unicode')
pb = xml.count('w:type="page"')

# Total = breaks + content overflow
print(f"Explicit page breaks: {pb}")
print(f"Content overflow pages: {total_pages}")
print(f"Estimated total pages: ~{pb + total_pages}")
